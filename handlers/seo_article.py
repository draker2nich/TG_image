from datetime import datetime
from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from aiogram.fsm.context import FSMContext
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.types import InlineKeyboardButton

from states.generation_states import SEOArticleStates
from keyboards.menus import cancel_kb, back_to_menu_kb, cancel_and_back_kb
from services.openai_service import openai_service
from services.google_service import google_service

router = Router()

# ID папки для статей на Google Drive
SEO_ARTICLES_FOLDER_ID = "1WDx-R5yz0nmTIHbLT4k_b5OzfTRwa8DH"


def confirm_outline_kb():
    """Кнопки для подтверждения структуры"""
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="✅ Подтвердить и генерировать", callback_data="seo:confirm_outline"))
    builder.row(InlineKeyboardButton(text="✏️ Редактировать структуру", callback_data="seo:edit_outline"))
    builder.row(InlineKeyboardButton(text="🔄 Перегенерировать структуру", callback_data="seo:regenerate_outline"))
    builder.row(InlineKeyboardButton(text="❌ Отмена", callback_data="cancel"))
    return builder.as_markup()


async def save_article_to_docx(article: str, outline: str, seo_title: str = "") -> bytes:
    """
    Сохраняет статью в формате DOCX.
    Структура берётся из outline, текст из article.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    import re
    
    doc = Document()
    
    # Парсим структуру из outline для получения заголовков
    outline_headers = []
    for line in outline.split('\n'):
        line = line.strip()
        if line.startswith('### '):
            outline_headers.append(('h3', line[4:]))
        elif line.startswith('## '):
            outline_headers.append(('h2', line[3:]))
        elif line.startswith('# '):
            outline_headers.append(('h1', line[2:]))
    
    # Парсим статью - разбиваем по заголовкам
    lines = article.split('\n')
    current_content = []
    current_header = None
    current_level = None
    
    # Добавляем H1 заголовок (SEO title или первый # заголовок)
    if seo_title:
        title = doc.add_heading(seo_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in lines:
        line_stripped = line.strip()
        
        # Определяем уровень заголовка
        if line_stripped.startswith('### '):
            # Сохраняем предыдущий контент
            if current_header and current_content:
                _add_content_to_doc(doc, current_content)
            current_header = line_stripped[4:]
            current_level = 3
            current_content = []
            doc.add_heading(current_header, level=3)
            
        elif line_stripped.startswith('## '):
            if current_header and current_content:
                _add_content_to_doc(doc, current_content)
            current_header = line_stripped[3:]
            current_level = 2
            current_content = []
            doc.add_heading(current_header, level=2)
            
        elif line_stripped.startswith('# '):
            if current_header and current_content:
                _add_content_to_doc(doc, current_content)
            # H1 уже добавлен как seo_title, пропускаем дубликат
            if not seo_title:
                current_header = line_stripped[2:]
                current_level = 1
                current_content = []
                doc.add_heading(current_header, level=1)
            else:
                current_header = line_stripped[2:]
                current_level = 1
                current_content = []
        else:
            # Обычный текст - добавляем к текущему контенту
            if line_stripped:
                current_content.append(line_stripped)
    
    # Добавляем последний блок контента
    if current_content:
        _add_content_to_doc(doc, current_content)
    
    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_content_to_doc(doc, content_lines: list):
    """Добавляет контент в документ"""
    import re
    from docx.shared import Pt
    
    for line in content_lines:
        # Пропускаем пустые строки
        if not line.strip():
            continue
        
        # Очищаем markdown форматирование
        clean_line = line
        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)
        clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
        clean_line = re.sub(r'__(.+?)__', r'\1', clean_line)
        
        # Обрабатываем списки
        if line.startswith('- ') or line.startswith('* '):
            clean_line = clean_line[2:].strip()
            para = doc.add_paragraph(clean_line, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            clean_line = re.sub(r'^\d+\.\s', '', clean_line).strip()
            para = doc.add_paragraph(clean_line, style='List Number')
        else:
            para = doc.add_paragraph(clean_line)
            para.paragraph_format.space_after = Pt(6)


async def upload_to_google(content: bytes, filename: str, title: str) -> str:
    """Загружает статью на Google Drive в папку для статей"""
    try:
        if not await google_service.initialize():
            return ""
        
        result = await google_service.upload_file_to_drive(
            file_content=content,
            file_name=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            folder_id=SEO_ARTICLES_FOLDER_ID
        )
        
        if result.success:
            await google_service.log_content(
                content_type="seo_article",
                title=title,
                status="uploaded",
                file_url=result.file_url or "",
                platform="blog"
            )
            return result.file_url or ""
        return ""
    except Exception as e:
        print(f"Upload error: {e}")
        return ""


@router.callback_query(F.data == "menu:seo")
async def start_seo_flow(callback: CallbackQuery, state: FSMContext):
    """Начало создания SEO-статьи"""
    if not openai_service.is_available():
        await callback.message.edit_text(
            "⚠️ OpenAI API не настроен.\nДобавьте OPENAI_API_KEY.",
            reply_markup=back_to_menu_kb()
        )
        await callback.answer()
        return
    
    await state.set_state(SEOArticleStates.waiting_topic)
    await callback.message.edit_text(
        "<b>📝 Создание SEO-статьи</b>\n\n"
        "Введите тему для статьи:",
        parse_mode="HTML",
        reply_markup=cancel_kb()
    )
    await callback.answer()


@router.message(SEOArticleStates.waiting_topic)
async def process_topic_and_generate_outline(message: Message, state: FSMContext):
    """Получение темы и генерация структуры"""
    topic = message.text.strip()
    await state.update_data(topic=topic)
    
    await message.answer("⏳ Генерирую структуру статьи...")
    
    try:
        # 1. Генерируем SEO-ключи
        seo_data = await openai_service.generate_seo_keywords(topic)
        keywords = [k.strip() for k in seo_data.get("keywords", "").split(",") if k.strip()]
        seo_title = seo_data.get("seo_title", topic)
        
        # 2. Генерируем структуру (строго в формате #, ##, ###)
        outline = await openai_service.generate_seo_outline(topic, keywords, seo_title)
        
        # Сохраняем данные
        await state.update_data(
            keywords=keywords,
            seo_title=seo_title,
            outline=outline
        )
        await state.set_state(SEOArticleStates.confirming_outline)
        
        # Показываем структуру
        await message.answer(
            f"📋 <b>Структура статьи</b>\n\n"
            f"<b>SEO-заголовок:</b> {seo_title}\n\n"
            f"<b>Ключи:</b> {', '.join(keywords[:5])}\n\n"
            f"<b>Структура:</b>\n<pre>{outline}</pre>\n\n"
            f"💡 Скопируйте структуру для редактирования если нужно изменить.",
            parse_mode="HTML",
            reply_markup=confirm_outline_kb()
        )
    except Exception as e:
        await message.answer(
            f"❌ Ошибка генерации структуры: {e}",
            reply_markup=back_to_menu_kb()
        )


@router.callback_query(SEOArticleStates.confirming_outline, F.data == "seo:confirm_outline")
async def confirm_and_generate_article(callback: CallbackQuery, state: FSMContext):
    """Подтверждение структуры и генерация статьи"""
    data = await state.get_data()
    
    await callback.message.edit_text("⏳ Генерирую статью... Это займёт 1-2 минуты.")
    await callback.answer()
    
    try:
        # Генерируем статью СТРОГО по структуре
        article = await openai_service.generate_seo_article_by_outline(
            outline=data['outline'],
            keywords=data['keywords'],
            seo_title=data['seo_title']
        )
        
        # Сохраняем в DOCX (структура + текст)
        docx_content = await save_article_to_docx(
            article=article,
            outline=data['outline'],
            seo_title=data['seo_title']
        )
        
        # Формируем имя файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = "".join(c if c.isalnum() or c in " -_" else "" for c in data['topic'])[:30]
        filename = f"SEO_{safe_topic.replace(' ', '_')}_{timestamp}.docx"
        
        # Загружаем на Google Drive
        google_url = await upload_to_google(docx_content, filename, data['seo_title'])
        
        # Отправляем пользователю
        file = BufferedInputFile(docx_content, filename=filename)
        
        await callback.message.answer_document(
            file,
            caption=(
                f"✅ <b>SEO-статья готова!</b>\n\n"
                f"📰 <b>Заголовок:</b> {data['seo_title']}\n"
                f"🔑 <b>Ключи:</b> {', '.join(data['keywords'][:5])}"
            ),
            parse_mode="HTML"
        )
        
        if google_url:
            await callback.message.answer(
                f"☁️ <a href='{google_url}'>Открыть на Google Drive</a>",
                parse_mode="HTML"
            )
        
        await state.clear()
        
        # Кнопка главного меню после генерации
        builder = InlineKeyboardBuilder()
        builder.row(InlineKeyboardButton(text="📝 Создать новую статью", callback_data="menu:seo"))
        builder.row(InlineKeyboardButton(text="⬅️ Главное меню", callback_data="menu:main"))
        
        await callback.message.answer(
            "✅ Статья успешно создана!\n\nВыберите следующее действие:",
            reply_markup=builder.as_markup()
        )
        
    except Exception as e:
        await callback.message.answer(
            f"❌ Ошибка генерации: {e}",
            reply_markup=back_to_menu_kb()
        )
        await state.clear()


@router.callback_query(SEOArticleStates.confirming_outline, F.data == "seo:edit_outline")
async def edit_outline(callback: CallbackQuery, state: FSMContext):
    """Редактирование структуры"""
    data = await state.get_data()
    
    await state.set_state(SEOArticleStates.editing_outline)
    await callback.message.edit_text(
        "✏️ <b>Редактирование структуры</b>\n\n"
        "Скопируйте структуру выше, отредактируйте и отправьте.\n\n"
        "<b>Формат заголовков:</b>\n"
        "• <code># Заголовок H1</code> (главный)\n"
        "• <code>## Заголовок H2</code> (раздел)\n"
        "• <code>### Заголовок H3</code> (подраздел)\n\n"
        f"<b>Текущая структура для копирования:</b>\n"
        f"<pre>{data['outline']}</pre>",
        parse_mode="HTML",
        reply_markup=cancel_and_back_kb("seo:back_outline")
    )
    await callback.answer()


@router.callback_query(SEOArticleStates.editing_outline, F.data == "seo:back_outline")
async def back_to_outline(callback: CallbackQuery, state: FSMContext):
    """Возврат к просмотру структуры"""
    data = await state.get_data()
    await state.set_state(SEOArticleStates.confirming_outline)
    
    await callback.message.edit_text(
        f"📋 <b>Структура статьи</b>\n\n"
        f"<b>SEO-заголовок:</b> {data['seo_title']}\n\n"
        f"<b>Ключи:</b> {', '.join(data['keywords'][:5])}\n\n"
        f"<b>Структура:</b>\n<pre>{data['outline']}</pre>",
        parse_mode="HTML",
        reply_markup=confirm_outline_kb()
    )
    await callback.answer()


@router.message(SEOArticleStates.editing_outline)
async def process_edited_outline(message: Message, state: FSMContext):
    """Обработка отредактированной структуры"""
    new_outline = message.text.strip()
    
    # Валидация структуры - должна содержать хотя бы один заголовок
    has_headers = any(
        line.strip().startswith('#') 
        for line in new_outline.split('\n')
    )
    
    if not has_headers:
        await message.answer(
            "⚠️ <b>Структура должна содержать заголовки!</b>\n\n"
            "Используйте формат:\n"
            "• <code># Заголовок H1</code>\n"
            "• <code>## Заголовок H2</code>\n"
            "• <code>### Заголовок H3</code>\n\n"
            "Попробуйте ещё раз:",
            parse_mode="HTML",
            reply_markup=cancel_and_back_kb("seo:back_outline")
        )
        return
    
    # Сохраняем новую структуру
    await state.update_data(outline=new_outline)
    await state.set_state(SEOArticleStates.confirming_outline)
    
    data = await state.get_data()
    
    await message.answer(
        f"✅ <b>Структура обновлена!</b>\n\n"
        f"<b>SEO-заголовок:</b> {data['seo_title']}\n\n"
        f"<b>Новая структура:</b>\n<pre>{new_outline}</pre>\n\n"
        f"💡 Статья будет сгенерирована СТРОГО по этой структуре.",
        parse_mode="HTML",
        reply_markup=confirm_outline_kb()
    )


@router.callback_query(SEOArticleStates.confirming_outline, F.data == "seo:regenerate_outline")
async def regenerate_outline(callback: CallbackQuery, state: FSMContext):
    """Перегенерация структуры"""
    data = await state.get_data()
    
    await callback.message.edit_text("⏳ Генерирую новую структуру...")
    await callback.answer()
    
    try:
        outline = await openai_service.generate_seo_outline(
            data['topic'], 
            data['keywords'], 
            data['seo_title']
        )
        
        await state.update_data(outline=outline)
        
        await callback.message.edit_text(
            f"📋 <b>Новая структура</b>\n\n"
            f"<b>SEO-заголовок:</b> {data['seo_title']}\n\n"
            f"<b>Ключи:</b> {', '.join(data['keywords'][:5])}\n\n"
            f"<b>Структура:</b>\n<pre>{outline}</pre>",
            parse_mode="HTML",
            reply_markup=confirm_outline_kb()
        )
    except Exception as e:
        await callback.message.edit_text(
            f"❌ Ошибка: {e}",
            reply_markup=back_to_menu_kb()
        )